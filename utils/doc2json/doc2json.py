from docx import Document
import os
import enum
import json
import re
import argparse
import hashlib

REGEX_TABLE_BACKGROUND = re.compile('w:fill=\"(\S*)\"')
REGEX_DEFINITION = re.compile('\s*(\d+)\s*\)\s*(.*)')
REGEX_CROSSWORD_ID = re.compile('\d+')

def is_cell_blocked(cell):
    match = REGEX_TABLE_BACKGROUND.search(cell._tc.xml)
    if match:
        result = match.group(1)
        assert(result in ["auto", "000000", "FFFFFF"])
        return (result != "auto" and result != "FFFFFF")
    
    return False

def process_docx_directory(directory_path, out_dir):
    seen = set()
    for filename in os.listdir(directory_path):
        if filename.endswith(".docx") and not "$" in filename:
            print(f"Parsing '{filename}'")
            file_path = os.path.join(directory_path, filename)
            cid = process_crossword(file_path, out_dir)
            if cid in seen:
                print(f"Error: ID {cid} already processed!")
                break
            seen.add(cid)

    missing = []
    if (len(seen)) == 0:
        return missing
    
    max_id = max(seen)
    for i in range(1, max_id + 1):
        if i not in seen:
            missing.append(i)

    return missing

class State(enum.Enum):
    READ_CROSSWORD_ID = "id"
    READ_AUTHOR = "author"
    SEARCH_FOR_DEFINITIONS = "search_defs"
    READ_DOWN_TITLE = "down_title"
    READ_DOWN_DEFINITIONS = "down"
    READ_ACROSS_TITLE = "across_title"
    READ_ACROSS_DEFINITIONS = "across"

def process_crossword(path, out_dir):
    document = Document(path)

    crossword = {
        State.READ_CROSSWORD_ID.value: 0,
        State.READ_AUTHOR.value: "",
        "dimensions": {
            "rows": 0,
            "columns": 0,
        },
        "grid": [],
        "definitions": {
            State.READ_DOWN_DEFINITIONS.value: {},
            State.READ_ACROSS_DEFINITIONS.value: {}
        }
    }

    state = State.READ_CROSSWORD_ID
    for para in document.paragraphs:
        line = para.text.strip()
        if line == "":
            continue

        if state == State.READ_CROSSWORD_ID:
            crossword[state.value] = int(REGEX_CROSSWORD_ID.findall(line)[0])
            state = State.READ_AUTHOR
        elif state == State.READ_AUTHOR:
            crossword[state.value] = line
            state = State.SEARCH_FOR_DEFINITIONS
        elif state == State.SEARCH_FOR_DEFINITIONS:
            if line == "מאונך":
                state = State.READ_DOWN_DEFINITIONS
            elif line == "מאוזן":
                state = State.READ_ACROSS_DEFINITIONS
            else:
                raise RuntimeError(f"Unexpected line: {line}")
        elif (state == State.READ_DOWN_DEFINITIONS) or (state == State.READ_ACROSS_DEFINITIONS):
            if line == "מאוזן":
                state = State.READ_ACROSS_DEFINITIONS
            elif line == "מאונך":
                state = State.READ_DOWN_DEFINITIONS
            else:
                match = re.search(REGEX_DEFINITION, line)
                if match:
                    number = int(match.group(1))
                    text = match.group(2)
                crossword["definitions"][state.value][number] = text
    assert( (state == State.READ_DOWN_DEFINITIONS) or (state == State.READ_ACROSS_DEFINITIONS) )
    for def_state in [State.READ_DOWN_DEFINITIONS.value, State.READ_ACROSS_DEFINITIONS.value]:
        assert(len(crossword["definitions"][def_state]) > 0)

    assert(len(document.tables) == 1)
    table = document.tables[0]
    crossword["dimensions"]["rows"] = len(table.rows)
    crossword["dimensions"]["columns"] = len(table.rows[0].cells)

    has_sol = False
    grid_sol = []
    for row in table.rows:
        row_list = []
        row_sol = []
        for cell in row.cells:
            #print(cell._tc.xml)
            if is_cell_blocked(cell):
                row_list.append('#')
                row_sol.append('#')
            else:
                text = "".join(p.text for p in cell.paragraphs).strip()
                hebrew_character = re.search(r'[א-ת]', text)
                if hebrew_character:
                    has_sol = True
                    row_sol.append(hebrew_character.group())
                number_match = re.search(r'\d+', text)
                row_list.append(number_match.group() if number_match is not None else "")
        crossword["grid"].append(row_list[::-1])
        grid_sol.append(row_sol[::-1])
        
        assert(len(row_list) == crossword["dimensions"]["columns"])

    if has_sol:
        crossword["sol_hash"] = hashlib.sha512("".join([item for sublist in grid_sol for item in sublist]).encode()).hexdigest()
        crossword["sol_grid"] = grid_sol
        try:
            crossword["solutions"] = extract_definition_solutions(crossword)
        except Exception as e:
            print(f"Error extracting definition solutions for crossword {crossword['id']}")
            print("This usually happens if the solution isn't complete.")
            print("For example: A missing character in the solution.")

    with open(os.path.join(out_dir, f"{crossword[State.READ_CROSSWORD_ID.value]}.json"), "w", encoding="utf8") as o:
        json.dump(crossword, o, indent=2)

    return crossword[State.READ_CROSSWORD_ID.value]

def get_def_sol(row, col, solution, progress):
    res = []
    dr, dc = progress

    def in_bounds(row, col):
        if row < 0 or row >= len(solution):
            return False
        if col < 0 or col >= len(solution[0]):
            return False
        return True

    while in_bounds(row, col) and (solution[row][col] != '#'):
        res.append(solution[row][col])
        row += dr
        col += dc
    
    return "".join(res)

def extract_definition_solutions(crossword_data):
    res = {}

    directions = {
        State.READ_DOWN_DEFINITIONS.value:   (1, 0), 
        State.READ_ACROSS_DEFINITIONS.value: (0, -1)
    }

    for k in directions:
        res[k] = {}

    for r in range(crossword_data["dimensions"]["rows"]):
        for c in range(crossword_data["dimensions"]["columns"]):
            if (crossword_data["grid"][r][c].isnumeric()):
                definition_number = crossword_data["grid"][r][c]
                for direction, progress in directions.items():
                    if int(definition_number) in crossword_data["definitions"][direction]:
                        res[direction][definition_number] = get_def_sol(r, c, 
                                                                        crossword_data["sol_grid"], 
                                                                        progress)
    return res

def main():
    parser = argparse.ArgumentParser(description="Convert Cryptic Crossword from docx to json")
    parser.add_argument("-i", "--input_path", required=True, help="Path to docx file or directory containing docx files")
    parser.add_argument("-o", "--out_dir", required=True, help="Output directory")
    parser.add_argument("-s", "--show_missing", action='store_true', help="List missing crosswords when parsing directory")
    args = parser.parse_args()

    input_path = args.input_path
    out_dir = args.out_dir

    if os.path.isfile(input_path):
        process_crossword(input_path, out_dir)
    elif os.path.isdir(input_path):
        missing = process_docx_directory(input_path, out_dir)
        if args.show_missing:
            print("The following crosswords are missing: ", missing)
    else:
        print("Invalid input path.")

if __name__ == "__main__":
    main()